"""
Report Generator Module
Supports PDF and DOCX export with proper formatting
"""

import io
import re
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Dict


def generate_pdf_report(markdown_content: str, title: str = "Research Report") -> bytes:
    """
    Generate PDF report from markdown content.
    
    Args:
        markdown_content: Markdown formatted report content
        title: Report title
        
    Returns:
        PDF file as bytes
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.platypus.tables import Table, TableStyle
        import markdown
        
        # Convert markdown to HTML first for basic parsing
        html_content = markdown.markdown(markdown_content)
        
        # Create PDF buffer
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        
        # Container for the 'Flowable' objects
        story = []
        
        # Styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1E88E5'),
            spaceAfter=30,
            alignment=1  # Center
        )
        
        heading2_style = ParagraphStyle(
            'Heading2Custom',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#1565C0'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        heading3_style = ParagraphStyle(
            'Heading3Custom',
            parent=styles['Heading3'],
            fontSize=14,
            textColor=colors.HexColor('#1976D2'),
            spaceAfter=10,
            spaceBefore=10
        )
        
        normal_style = ParagraphStyle(
            'NormalCustom',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
            spaceAfter=6
        )
        
        # Add title
        story.append(Paragraph(title.replace('_', ' ').title(), title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Add generation timestamp
        timestamp = f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        story.append(Paragraph(f"<i>{timestamp}</i>", normal_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Parse and add content sections
        lines = markdown_content.split('\n')
        current_list = []
        in_list = False
        
        for line in lines:
            stripped = line.strip()
            
            if not stripped:
                if current_list and in_list:
                    # Add accumulated list items
                    list_data = [[Paragraph(item, normal_style)] for item in current_list]
                    list_table = Table(list_data, colWidths=[doc.width])
                    list_table.setStyle(TableStyle([
                        ('LEFTPADDING', (0, 0), (-1, -1), 12),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 12),
                        ('TOPPADDING', (0, 0), (-1, -1), 2),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                    ]))
                    story.append(list_table)
                    current_list = []
                    in_list = False
                story.append(Spacer(1, 0.1*inch))
                continue
            
            # Headers
            if stripped.startswith('# '):
                story.append(Paragraph(stripped[2:], title_style))
                story.append(Spacer(1, 0.2*inch))
            elif stripped.startswith('## '):
                story.append(Paragraph(stripped[3:], heading2_style))
                story.append(Spacer(1, 0.15*inch))
            elif stripped.startswith('### '):
                story.append(Paragraph(stripped[4:], heading3_style))
                story.append(Spacer(1, 0.1*inch))
            
            # List items
            elif stripped.startswith('- ') or stripped.startswith('* ') or stripped.startswith('• '):
                in_list = True
                current_list.append(stripped[2:])
            
            # Regular paragraphs
            else:
                if current_list and in_list:
                    # Flush list before adding paragraph
                    list_data = [[Paragraph(item, normal_style)] for item in current_list]
                    list_table = Table(list_data, colWidths=[doc.width])
                    list_table.setStyle(TableStyle([
                        ('LEFTPADDING', (0, 0), (-1, -1), 12),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 12),
                        ('TOPPADDING', (0, 0), (-1, -1), 2),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                    ]))
                    story.append(list_table)
                    current_list = []
                    in_list = False
                
                # Handle bold and italic - need to be careful with order
                formatted = stripped
                
                # First, escape HTML special characters
                formatted = formatted.replace('&', '&amp;')
                formatted = formatted.replace('<', '&lt;')
                formatted = formatted.replace('>', '&gt;')
                
                # Now convert markdown to ReportLab-compatible HTML
                # Process bold first (longer patterns)
                formatted = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', formatted)
                formatted = re.sub(r'__(.+?)__', r'<b>\1</b>', formatted)
                # Then italic
                formatted = re.sub(r'(?<!\w)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'<i>\1</i>', formatted)
                formatted = re.sub(r'(?<!\w)_(?!_)(.+?)(?<!_)_(?!_)', r'<i>\1</i>', formatted)
                
                try:
                    story.append(Paragraph(formatted, normal_style))
                except Exception as e:
                    # Fallback: use plain text if formatting fails
                    # Strip all markdown and HTML
                    plain_text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
                    plain_text = re.sub(r'__(.+?)__', r'\1', plain_text)
                    plain_text = re.sub(r'\*(.+?)\*', r'\1', plain_text)
                    plain_text = re.sub(r'_(.+?)_', r'\1', plain_text)
                    # Escape HTML in plain text
                    plain_text = plain_text.replace('&', '&amp;')
                    plain_text = plain_text.replace('<', '&lt;')
                    plain_text = plain_text.replace('>', '&gt;')
                    story.append(Paragraph(plain_text, normal_style))
        
        # Add any remaining list
        if current_list:
            list_data = [[Paragraph(item, normal_style)] for item in current_list]
            list_table = Table(list_data, colWidths=[doc.width])
            list_table.setStyle(TableStyle([
                ('LEFTPADDING', (0, 0), (-1, -1), 12),
                ('RIGHTPADDING', (0, 0), (-1, -1), 12),
                ('TOPPADDING', (0, 0), (-1, -1), 2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ]))
            story.append(list_table)
        
        # Build PDF
        doc.build(story)
        
        # Get PDF data
        pdf_data = buffer.getvalue()
        buffer.close()
        
        return pdf_data
        
    except ImportError:
        # Fallback: return markdown as bytes if reportlab not available
        raise ImportError(
            "reportlab and markdown packages required for PDF generation. "
            "Install with: pip install reportlab markdown"
        )


def generate_docx_report(markdown_content: str, title: str = "Research Report") -> bytes:
    """
    Generate DOCX report from markdown content.
    
    Args:
        markdown_content: Markdown formatted report content
        title: Report title
        
    Returns:
        DOCX file as bytes
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        
        # Create document
        doc = Document()
        
        # Title
        title_para = doc.add_heading(title.replace('_', ' ').title(), 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Timestamp
        timestamp = doc.add_paragraph()
        timestamp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = timestamp.add_run(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        run.italic = True
        
        doc.add_paragraph()
        
        # Parse markdown content
        lines = markdown_content.split('\n')
        current_list = None
        
        for line in lines:
            stripped = line.strip()
            
            if not stripped:
                if current_list is not None:
                    current_list = None
                continue
            
            # Headers
            if stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=1)
                current_list = None
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
                current_list = None
            elif stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
                current_list = None
            
            # List items
            elif stripped.startswith('- ') or stripped.startswith('* ') or stripped.startswith('• '):
                if current_list is None:
                    current_list = doc.add_paragraph(style='List Bullet')
                else:
                    current_list = doc.add_paragraph(style='List Bullet')
                
                item_text = stripped[2:]
                # Handle inline formatting
                item_text = item_text.replace('**', '').replace('__', '')
                item_text = item_text.replace('*', '').replace('_', '')
                current_list.add_run(item_text)
            
            # Regular paragraphs
            else:
                current_list = None
                para = doc.add_paragraph()
                
                # Simple formatting detection
                text = stripped
                # Remove markdown syntax for now (could be enhanced)
                text = text.replace('**', '').replace('__', '')
                text = text.replace('*', '').replace('_', '')
                text = text.replace('`', '')
                
                para.add_run(text)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        docx_data = buffer.read()
        buffer.close()
        
        return docx_data
        
    except ImportError:
        raise ImportError(
            "python-docx package required for DOCX generation. "
            "Install with: pip install python-docx"
        )


def get_report_download_button(
    st,
    content: str,
    filename_base: str = "research_report",
    title: str = "Research Report"
):
    """
    Create download buttons for PDF and DOCX formats in Streamlit.
    
    Args:
        st: Streamlit module
        content: Markdown report content
        filename_base: Base filename (without extension)
        title: Report title for metadata
    """
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        try:
            pdf_data = generate_pdf_report(content, title)
            st.download_button(
                label="📄 Download PDF",
                data=pdf_data,
                file_name=f"{filename_base}_{timestamp}.pdf",
                mime="application/pdf",
                use_container_width=True
            )
        except ImportError as e:
            st.warning(f"PDF export unavailable: {str(e)}")
    
    with col2:
        try:
            docx_data = generate_docx_report(content, title)
            st.download_button(
                label="📝 Download DOCX",
                data=docx_data,
                file_name=f"{filename_base}_{timestamp}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        except ImportError as e:
            st.warning(f"DOCX export unavailable: {str(e)}")
    
    with col3:
        st.download_button(
            label="📋 Download Markdown",
            data=content,
            file_name=f"{filename_base}_{timestamp}.md",
            mime="text/markdown",
            use_container_width=True
        )


class ReportGenerator:
    """
    Unified report generator supporting multiple formats.
    """
    
    def __init__(self, title: str = "Research Report"):
        self.title = title
        self.content = ""
    
    def set_content(self, markdown_content: str):
        """Set report content."""
        self.content = markdown_content
    
    def generate_pdf(self) -> bytes:
        """Generate PDF format."""
        return generate_pdf_report(self.content, self.title)
    
    def generate_docx(self) -> bytes:
        """Generate DOCX format."""
        return generate_docx_report(self.content, self.title)
    
    def generate_all_formats(self) -> Dict[str, bytes]:
        """Generate all supported formats."""
        results = {}
        
        try:
            results['pdf'] = self.generate_pdf()
        except Exception as e:
            results['pdf_error'] = str(e)
        
        try:
            results['docx'] = self.generate_docx()
        except Exception as e:
            results['docx_error'] = str(e)
        
        results['markdown'] = self.content.encode('utf-8')
        
        return results
